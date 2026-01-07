from flask import Flask, request, jsonify, send_file
import os
import json
import re
import io
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload, MediaIoBaseDownload
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict

app = Flask(__name__)

# --- CONFIGURAÇÃO DE CREDENCIAIS ---
# Na Vercel, você deve criar uma variável de ambiente chamada 'GOOGLE_CREDENTIALS'
# contendo o CONTEÚDO do arquivo JSON da sua Service Account.
def get_services():
    creds_json = os.environ.get('GOOGLE_CREDENTIALS')
    if not creds_json:
        raise Exception("Credenciais do Google não configuradas (Variável GOOGLE_CREDENTIALS).")
    
    info = json.loads(creds_json)
    creds = service_account.Credentials.from_service_account_info(
        info, scopes=['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/spreadsheets']
    )
    return build('sheets', 'v4', credentials=creds), build('drive', 'v3', credentials=creds)

def extrair_id(link):
    m = re.search(r"/d/([a-zA-Z0-9-]+)", link)
    return m.group(1) if m else link.strip()

@app.route('/api/obter_colunas', methods=['POST'])
def obter_colunas():
    try:
        data = request.json
        link = data.get('link')
        aba = data.get('aba')
        
        sheets_service, _ = get_services()
        sid = extrair_id(link)
        
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=sid, range=f"'{aba}'!A2:Z2").execute()
        valores = result.get('values', [[]])
        cabecalho = valores[0]
        
        # Retorna lista de objetos para o frontend
        colunas = [{"id": i, "nome": n.strip()} for i, n in enumerate(cabecalho) if n.strip()]
        return jsonify(colunas)
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/processar', methods=['POST'])
def processar():
    try:
        data = request.json
        link = data.get('link')
        aba = data.get('aba')
        letra_escola = data.get('letra_escola', 'A')
        filtro_exclusao = data.get('filtro_exclusao', '')
        indices_excluir = data.get('indices_excluir', []) # Lista de inteiros
        formato = data.get('formato', 'DOCS')

        sheets_service, drive_service = get_services()
        sid = extrair_id(link)

        # 1. Busca Dados
        result = sheets_service.spreadsheets().values().get(spreadsheetId=sid, range=f"'{aba}'!A:Z").execute()
        linhas_todas = result.get('values', [])
        
        if len(linhas_todas) < 2:
            return jsonify({"erro": "Planilha sem dados na Linha 2."}), 400

        cabecalho_original = linhas_todas[1]
        idx_esc = ord(letra_escola.upper()) - ord('A')

        # 2. Define Colunas (Lógica original)
        colunas_que_ficam = []
        for i in range(len(cabecalho_original)):
            if i not in indices_excluir and cabecalho_original[i].strip():
                colunas_que_ficam.append((i, cabecalho_original[i].strip()))

        # 3. Agrupamento
        grupos = defaultdict(list)
        for linha in linhas_todas[2:]:
            if not linha: continue
            
            # Filtro "Não Conter"
            texto_unido = " ".join(map(str, linha)).lower()
            if filtro_exclusao and filtro_exclusao.lower() in texto_unido:
                continue

            escola = str(linha[idx_esc]).strip() if idx_esc < len(linha) else "GERAL"

            dados_celulas = []
            for idx_col, _ in colunas_que_ficam:
                valor = str(linha[idx_col] if idx_col < len(linha) else "").strip()
                dados_celulas.append(valor)
            
            grupos[escola].append(dados_celulas)

        # 4. Geração DOCX
        doc = Document()
        p = doc.add_paragraph("PORTARIA N° XXXX/2026-GS/SEDUC")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = doc.add_paragraph("O(A) SECRETÁRIO(A) DE ESTADO DE EDUCAÇÃO DO PARÁ, no uso de suas atribuições...")
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for escola, lista_dados in grupos.items():
            doc.add_paragraph("").add_run(f"ANEXO - {escola}").bold = True
            
            tabela = doc.add_table(rows=1, cols=len(colunas_que_ficam))
            tabela.style = 'Table Grid'
            
            # Cabeçalho
            for j, (_, nome_col) in enumerate(colunas_que_ficam):
                celula = tabela.cell(0, j)
                celula.text = nome_col
                for p in celula.paragraphs:
                    for r in p.runs: r.font.name = "Arial"; r.font.size = Pt(6); r.bold = True
            
            # Dados
            for registro in lista_dados:
                row_cells = tabela.add_row().cells
                for j, valor in enumerate(registro):
                    row_cells[j].text = valor
                    for p in row_cells[j].paragraphs:
                        for r in p.runs: r.font.name = "Arial"; r.font.size = Pt(6)
            doc.add_paragraph("")

        # 5. Salva em Buffer
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        # 6. Upload para Drive (Necessário para converter ou gerar link)
        nome_arquivo = f"PORTARIA_GERADA_{escola[:10]}"
        media = MediaIoBaseUpload(stream, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file_metadata = {'name': nome_arquivo, 'mimeType': 'application/vnd.google-apps.document'}
        
        gdoc = drive_service.files().create(body=file_metadata, media_body=media, fields='id, webViewLink').execute()
        fid = gdoc.get('id')

        # Se for PDF
        if formato == 'PDF':
            req = drive_service.files().export_media(fileId=fid, mimeType='application/pdf')
            pdf_fh = io.BytesIO()
            downloader = MediaIoBaseDownload(pdf_fh, req)
            done = False
            while not done: _, done = downloader.next_chunk()
            
            # Opcional: Deletar arquivo temporário do Drive para não lotar
            try:
                drive_service.files().delete(fileId=fid).execute()
            except: pass

            pdf_fh.seek(0)
            return send_file(
                pdf_fh,
                as_attachment=True,
                download_name=f"{nome_arquivo}.pdf",
                mimetype='application/pdf'
            )

        # Se for DOCS (Retorna o Link)
        return jsonify({
            "status": "sucesso",
            "link": gdoc.get('webViewLink'),
            "mensagem": "Arquivo criado com sucesso!"
        })

    except Exception as e:
        return jsonify({"erro": str(e)}), 500

# Necessário para Vercel
if __name__ == '__main__':
    app.run(debug=True)
