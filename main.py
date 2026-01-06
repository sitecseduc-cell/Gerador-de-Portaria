import re, io
from google.colab import auth, output, files
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict

# --- AUTENTICAÇÃO ---
auth.authenticate_user()
sheets_service = build('sheets', 'v4')
drive_service = build('drive', 'v3')

def extrair_id(link):
    m = re.search(r"/d/([a-zA-Z0-9-]+)", link)
    return m.group(1) if m else link.strip()

def obter_colunas(link, aba):
    try:
        sid = extrair_id(link)
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=sid, range=f"'{aba}'!A2:Z2").execute()
        valores = result.get('values', [[]])
        cabecalho = valores[0]
        return [f"{i}|{n.strip()}" for i, n in enumerate(cabecalho) if n.strip()]
    except Exception as e: return [f"Erro: {str(e)}"]

def processar_documento(link, aba, letra_escola, exc, excluir_str, formato):
    try:
        sid = extrair_id(link)
        result = sheets_service.spreadsheets().values().get(spreadsheetId=sid, range=f"'{aba}'!A:Z").execute()
        linhas_todas = result.get('values', [])
        
        if len(linhas_todas) < 2: return "❌ Erro: Planilha sem dados."

        cabecalho_original = linhas_todas[1]
        idx_esc = ord(letra_escola.upper()) - ord('A')
        indices_para_remover = [int(i) for i in excluir_str.split(',') if i.strip()]
        
        colunas_que_ficam = []
        for i in range(len(cabecalho_original)):
            if i not in indices_para_remover and cabecalho_original[i].strip():
                colunas_que_ficam.append((i, cabecalho_original[i].strip()))

        grupos = defaultdict(list)
        for linha in linhas_todas[2:]:
            if not linha: continue
            texto_unido = " ".join(map(str, linha)).lower()
            if exc and exc.lower() in texto_unido: continue
            
            escola = str(linha[idx_esc]).strip() if idx_esc < len(linha) else "GERAL"
            dados_celulas = [str(linha[idx] if idx < len(linha) else "").strip() for idx, _ in colunas_que_ficam]
            grupos[escola].append(dados_celulas)

        doc = Document()
        # (Cabeçalho da Portaria)
        p = doc.add_paragraph("PORTARIA N° XXXX/2026-GS/SEDUC")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("O(A) SECRETÁRIO(A) DE ESTADO DE EDUCAÇÃO DO PARÁ...").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for escola, lista_dados in grupos.items():
            doc.add_paragraph("").add_run(f"ANEXO - {escola}").bold = True
            tabela = doc.add_table(rows=1, cols=len(colunas_que_ficam))
            tabela.style = 'Table Grid'
            
            for j, (_, nome_col) in enumerate(colunas_que_ficam):
                celula = tabela.cell(0, j)
                celula.text = nome_col
                for p in celula.paragraphs:
                    for r in p.runs: r.font.name = "Arial"; r.font.size = Pt(6); r.bold = True

            for registro in lista_dados:
                row_cells = tabela.add_row().cells
                for j, valor in enumerate(registro):
                    row_cells[j].text = valor
                    for p in row_cells[j].paragraphs:
                        for r in p.runs: r.font.name = "Arial"; r.font.size = Pt(6)

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        media = MediaIoBaseUpload(stream, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        gdoc = drive_service.files().create(body={'name': 'PORTARIA_GERADA', 'mimeType': 'application/vnd.google-apps.document'}, media_body=media).execute()
        
        if formato == 'PDF':
            req = drive_service.files().export_media(fileId=gdoc['id'], mimeType='application/pdf')
            pdf_fh = io.BytesIO()
            downloader = MediaIoBaseDownload(pdf_fh, req)
            done = False
            while not done: _, done = downloader.next_chunk()
            with open("PORTARIA.pdf", "wb") as f: f.write(pdf_fh.getvalue())
            drive_service.files().delete(fileId=gdoc['id']).execute()
            files.download("PORTARIA.pdf")
            return "✅ PDF baixado com sucesso!"

        return f"✅ Google Docs criado! ID: {gdoc['id']}"
    except Exception as e: return f"❌ Erro: {str(e)}"

# Registro para o Colab reconhecer as funções
output.register_callback('notebook.obter_colunas', obter_colunas)
output.register_callback('notebook.processar_documento', processar_documento)
